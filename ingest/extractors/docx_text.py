"""Extract plain text from DOCX bytes using python-docx (OSS)."""

from __future__ import annotations

import io


def extract_docx_text(data: bytes) -> str:
    if not data or not data.strip():
        return ""
    try:
        import docx
    except ImportError as exc:
        raise ValueError("DOCX extraction requires the 'python-docx' package.") from exc

    document = docx.Document(io.BytesIO(data))
    parts: list[str] = []
    for para in document.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in document.tables:
        for row in table.rows:
            cells = " | ".join((c.text or "").strip() for c in row.cells)
            if cells.strip():
                parts.append(cells)
    return "\n".join(parts).strip()
